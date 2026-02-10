from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import os
import re

router = APIRouter()

class ScenarioExportRequest(BaseModel):
    title: str
    scenario_type: str
    content: str
    exam_type: str = "E4"
    student_name: str = ""

def extract_field_value(content: str, field_name: str) -> str:
    """Extract a field value from the generated content"""
    patterns = [
        rf"{field_name}\s*:?\s*([^\n]+)",
        rf"#{1,3}\s*{field_name}\s*\n([^\n#]+)",
        rf"\*\*{field_name}\*\*\s*:?\s*([^\n]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return ""

def parse_scenario_content(content: str) -> dict:
    """Parse the generated scenario content into structured fields"""
    fields = {
        'objet': '',
        'date_duree': '',
        'lieu': '',
        'delimitation': '',
        'acteurs': '',
        'historique': '',
        'objectifs': '',
        'informations': '',
        'contraintes': '',
        'identite_client': '',
        'relation_entreprise': '',
        'date_rencontre': '',
        'motivations': '',
        'freins': '',
        'objections': ''
    }
    
    # Try to extract structured information
    lines = content.split('\n')
    current_section = None
    section_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if it's a section header
        if line.startswith('#'):
            # Save previous section
            if current_section and section_content:
                content_text = ' '.join(section_content)
                
                # Map sections to fields
                section_lower = current_section.lower()
                if 'objet' in section_lower or 'contexte' in section_lower:
                    fields['objet'] = content_text
                elif 'date' in section_lower or 'durée' in section_lower:
                    fields['date_duree'] = content_text
                elif 'lieu' in section_lower:
                    fields['lieu'] = content_text
                elif 'acteur' in section_lower or 'client' in section_lower or 'identité' in section_lower:
                    fields['acteurs'] = content_text
                    fields['identite_client'] = content_text
                elif 'historique' in section_lower or 'relation' in section_lower:
                    fields['historique'] = content_text
                    fields['relation_entreprise'] = content_text
                elif 'objectif' in section_lower:
                    fields['objectifs'] = content_text
                elif 'information' in section_lower or 'donnée' in section_lower:
                    fields['informations'] = content_text
                elif 'contrainte' in section_lower:
                    fields['contraintes'] = content_text
                elif 'motivation' in section_lower:
                    fields['motivations'] = content_text
                elif 'frein' in section_lower or 'obstacle' in section_lower:
                    fields['freins'] = content_text
                elif 'objection' in section_lower:
                    fields['objections'] = content_text
            
            # Start new section
            current_section = line.replace('#', '').strip()
            section_content = []
        else:
            section_content.append(line)
    
    # Save last section
    if current_section and section_content:
        content_text = ' '.join(section_content)
        section_lower = current_section.lower()
        if 'objet' in section_lower:
            fields['objet'] = content_text
        elif 'objectif' in section_lower:
            fields['objectifs'] = content_text
    
    # If no structured data found, use the title and full content
    if not any(fields.values()):
        fields['objet'] = content[:200] if len(content) > 200 else content
        fields['objectifs'] = "Réaliser la simulation selon le scénario fourni"
    
    return fields

@router.post("/export-scenario/docx")
async def export_scenario_docx(request: ScenarioExportRequest):
    """Export scenario as Word using official template"""
    
    # Load template
    template_path = os.path.join(os.path.dirname(__file__), '..', '..', 'templates', 'PARAMETRES MODIFIES 2024 vierge- Copie copie.docx')
    
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Template not found")
    
    doc = Document(template_path)
    
    # Parse the generated content
    fields = parse_scenario_content(request.content)
    
    # Update the student name in paragraphs
    for para in doc.paragraphs:
        if 'nom du CANDIDAT' in para.text:
            para.text = para.text.replace('nom du CANDIDAT :', request.student_name or 'CANDIDAT')
            # Restore formatting
            for run in para.runs:
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
                run.font.bold = True
        
        # Update checkboxes based on scenario type
        if 'Négociation Vente' in para.text:
            for run in para.runs:
                if run.font.name == 'Wingdings':
                    is_negociation = 'negociation' in request.scenario_type.lower()
                    run.text = 'þ' if is_negociation else 'o'
        elif 'Evènement commercial' in para.text or 'Événement commercial' in para.text:
            for run in para.runs:
                if run.font.name == 'Wingdings':
                    is_event = 'evenement' in request.scenario_type.lower() or 'event' in request.scenario_type.lower()
                    run.text = 'þ' if is_event else 'o'
    
    # Fill Table 1 (Candidat) - 10 rows
    if len(doc.tables) >= 1:
        table = doc.tables[0]
        
        # Row 2: Objet de l'activité - USE ONLY THE TITLE, NOT FULL CONTENT
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            # Use the title (which is the scenario name), not the full content
            table.rows[1].cells[1].text = request.title[:150]  # Limit to 150 chars
        
        # Row 3: Date(s) et durée
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            table.rows[2].cells[1].text = fields['date_duree'][:100] if fields['date_duree'] else "À définir"
        
        # Row 4: Lieu
        if len(table.rows) > 3 and len(table.rows[3].cells) > 1:
            table.rows[3].cells[1].text = fields['lieu'][:100] if fields['lieu'] else "À définir"
        
        # Row 5: Délimitation de Séquence(s)
        if len(table.rows) > 4 and len(table.rows[4].cells) > 1:
            table.rows[4].cells[1].text = fields['delimitation'][:150] if fields['delimitation'] else "Durée: 30-40 minutes"
        
        # Row 6: Acteur(s) concernés
        if len(table.rows) > 5 and len(table.rows[5].cells) > 1:
            acteurs_text = fields['acteurs'] or fields['identite_client']
            table.rows[5].cells[1].text = acteurs_text[:200] if acteurs_text else "Client/Prospect"
        
        # Row 7: Historique de la relation
        if len(table.rows) > 6 and len(table.rows[6].cells) > 1:
            historique_text = fields['historique'] or fields['relation_entreprise']
            table.rows[6].cells[1].text = historique_text[:250] if historique_text else "Première prise de contact"
        
        # Row 8: Objectifs de la simulation
        if len(table.rows) > 7 and len(table.rows[7].cells) > 1:
            table.rows[7].cells[1].text = fields['objectifs'][:200] if fields['objectifs'] else "Réaliser la simulation selon le scénario"
        
        # Row 9: Informations à exploiter
        if len(table.rows) > 8 and len(table.rows[8].cells) > 1:
            # Extract key information, not the full content
            info_text = fields['informations'] if fields['informations'] else "Voir le scénario détaillé ci-dessous"
            table.rows[8].cells[1].text = info_text[:300]
        
        # Row 10: Contrainte(s)
        if len(table.rows) > 9 and len(table.rows[9].cells) > 1:
            table.rows[9].cells[1].text = fields['contraintes'][:200] if fields['contraintes'] else "Respecter le cadre professionnel"
    
    # Fill Table 2 (Jury) - 13 rows
    if len(doc.tables) >= 2:
        table = doc.tables[1]
        
        # Row 2: Objet de l'activité - USE ONLY THE TITLE
        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
            table.rows[1].cells[1].text = request.title[:150]
        
        # Row 3: Identité
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            identite_text = fields['identite_client'] or fields['acteurs']
            table.rows[2].cells[1].text = identite_text[:200] if identite_text else "Profil à définir"
        
        # Row 4: Relation à l'entreprise
        if len(table.rows) > 3 and len(table.rows[3].cells) > 1:
            table.rows[3].cells[1].text = fields['relation_entreprise'][:150] if fields['relation_entreprise'] else "Client potentiel"
        
        # Row 5: Date de la rencontre
        if len(table.rows) > 4 and len(table.rows[4].cells) > 1:
            date_text = fields['date_rencontre'] or fields['date_duree']
            table.rows[4].cells[1].text = date_text[:100] if date_text else "À définir"
        
        # Row 6: Lieu
        if len(table.rows) > 5 and len(table.rows[5].cells) > 1:
            table.rows[5].cells[1].text = fields['lieu'][:100] if fields['lieu'] else "À définir"
        
        # Row 7: Historique de la relation
        if len(table.rows) > 6 and len(table.rows[6].cells) > 1:
            table.rows[6].cells[1].text = fields['historique'][:250] if fields['historique'] else "Première prise de contact"
        
        # Row 8: Objectifs de la simulation
        if len(table.rows) > 7 and len(table.rows[7].cells) > 1:
            table.rows[7].cells[1].text = fields['objectifs'][:200] if fields['objectifs'] else "Jouer le rôle selon le profil défini"
        
        # Row 9: Délimitation de Séquence(s)
        if len(table.rows) > 8 and len(table.rows[8].cells) > 1:
            table.rows[8].cells[1].text = fields['delimitation'][:150] if fields['delimitation'] else "30-40 minutes"
        
        # Row 10: Motivations
        if len(table.rows) > 9 and len(table.rows[9].cells) > 1:
            table.rows[9].cells[1].text = fields['motivations'][:200] if fields['motivations'] else "À définir selon le profil"
        
        # Row 11: Freins
        if len(table.rows) > 10 and len(table.rows[10].cells) > 1:
            table.rows[10].cells[1].text = fields['freins'][:200] if fields['freins'] else "À définir selon le profil"
        
        # Row 12: Contrainte(s)
        if len(table.rows) > 11 and len(table.rows[11].cells) > 1:
            table.rows[11].cells[1].text = fields['contraintes'][:200] if fields['contraintes'] else "Respecter le cadre professionnel"
        
        # Row 13: Objections
        if len(table.rows) > 12 and len(table.rows[12].cells) > 1:
            table.rows[12].cells[1].text = fields['objections'][:200] if fields['objections'] else "À définir selon le contexte"
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Sujet_{request.exam_type}_{request.student_name or 'CANDIDAT'}.docx"}
    )

@router.post("/export-scenario/pdf")
async def export_scenario_pdf(request: ScenarioExportRequest):
    """Export scenario as PDF matching the Word template structure"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=2.5*cm, leftMargin=2.5*cm,
                           topMargin=2.5*cm, bottomMargin=2.5*cm)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=10,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        spaceAfter=6
    )
    
    header_style = ParagraphStyle(
        'HeaderStyle',
        fontSize=9,
        fontName='Helvetica-Bold',
        alignment=TA_LEFT
    )
    
    cell_style = ParagraphStyle(
        'CellStyle',
        fontSize=9,
        fontName='Helvetica',
        alignment=TA_LEFT
    )
    
    # Parse content
    fields = parse_scenario_content(request.content)
    
    # Header
    story.append(Paragraph("BTS Négociation et Digitalisation de la Relation Client", title_style))
    story.append(Paragraph("Session 2024", title_style))
    exam_num = request.exam_type[-1] if request.exam_type else '4'
    story.append(Paragraph(f"E{exam_num} – RELATION CLIENT et NEGOCIATION VENTE", title_style))
    story.append(Spacer(1, 0.5*cm))
    
    # Fiche sujet
    story.append(Paragraph(f"FICHE SUJET – {request.student_name or 'CANDIDAT'}", title_style))
    story.append(Spacer(1, 0.3*cm))
    
    # Checkboxes
    is_negociation = 'negociation' in request.scenario_type.lower()
    is_event = 'evenement' in request.scenario_type.lower() or 'event' in request.scenario_type.lower()
    checkbox_neg = "☑" if is_negociation else "☐"
    checkbox_evt = "☑" if is_event else "☐"
    
    story.append(Paragraph(f"{checkbox_neg} Négociation Vente et Accompagnement de la Relation Client", cell_style))
    story.append(Paragraph(f"{checkbox_evt} Organisation et Animation d'un Evènement commercial", cell_style))
    story.append(Spacer(1, 0.5*cm))
    
    # TABLE 1: FICHE CANDIDAT
    story.append(Paragraph("FICHE CANDIDAT", title_style))
    story.append(Spacer(1, 0.2*cm))
    
    table1_data = [
        [Paragraph("Objet de l'activité", header_style), 
         Paragraph(request.title[:150], cell_style)],
        [Paragraph("Date(s) et durée", header_style), 
         Paragraph(fields['date_duree'][:100] if fields['date_duree'] else "À définir", cell_style)],
        [Paragraph("Lieu", header_style), 
         Paragraph(fields['lieu'][:100] if fields['lieu'] else "À définir", cell_style)],
        [Paragraph("Délimitation de Séquence(s)", header_style), 
         Paragraph(fields['delimitation'][:150] if fields['delimitation'] else "Durée: 30-40 minutes", cell_style)],
        [Paragraph("Acteur(s) concernés", header_style), 
         Paragraph((fields['acteurs'] or fields['identite_client'] or "Client/Prospect")[:200], cell_style)],
        [Paragraph("Historique de la relation", header_style), 
         Paragraph((fields['historique'] or fields['relation_entreprise'] or "Première prise de contact")[:250], cell_style)],
        [Paragraph("Objectifs de la simulation", header_style), 
         Paragraph((fields['objectifs'] or "Réaliser la simulation selon le scénario")[:200], cell_style)],
        [Paragraph("Informations à exploiter", header_style), 
         Paragraph((fields['informations'] or "Voir le scénario détaillé")[:300], cell_style)],
        [Paragraph("Contrainte(s)", header_style), 
         Paragraph((fields['contraintes'] or "Respecter le cadre professionnel")[:200], cell_style)],
    ]
    
    table1 = Table(table1_data, colWidths=[5*cm, 11*cm])
    table1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    
    story.append(table1)
    story.append(PageBreak())
    
    # TABLE 2: FICHE JURY
    story.append(Paragraph("FICHE JURY", title_style))
    story.append(Spacer(1, 0.2*cm))
    
    table2_data = [
        [Paragraph("Objet de l'activité", header_style), 
         Paragraph(request.title[:150], cell_style)],
        [Paragraph("Identité", header_style), 
         Paragraph((fields['identite_client'] or fields['acteurs'] or "Profil à définir")[:200], cell_style)],
        [Paragraph("Relation à l'entreprise", header_style), 
         Paragraph((fields['relation_entreprise'] or "Client potentiel")[:150], cell_style)],
        [Paragraph("Date de la rencontre", header_style), 
         Paragraph((fields['date_rencontre'] or fields['date_duree'] or "À définir")[:100], cell_style)],
        [Paragraph("Lieu", header_style), 
         Paragraph((fields['lieu'] or "À définir")[:100], cell_style)],
        [Paragraph("Historique de la relation", header_style), 
         Paragraph((fields['historique'] or "Première prise de contact")[:250], cell_style)],
        [Paragraph("Objectifs de la simulation", header_style), 
         Paragraph((fields['objectifs'] or "Jouer le rôle selon le profil défini")[:200], cell_style)],
        [Paragraph("Délimitation de Séquence(s)", header_style), 
         Paragraph((fields['delimitation'] or "30-40 minutes")[:150], cell_style)],
        [Paragraph("Motivations", header_style), 
         Paragraph((fields['motivations'] or "À définir selon le profil")[:200], cell_style)],
        [Paragraph("Freins", header_style), 
         Paragraph((fields['freins'] or "À définir selon le profil")[:200], cell_style)],
        [Paragraph("Contrainte(s)", header_style), 
         Paragraph((fields['contraintes'] or "Respecter le cadre professionnel")[:200], cell_style)],
        [Paragraph("Objections", header_style), 
         Paragraph((fields['objections'] or "À définir selon le contexte")[:200], cell_style)],
    ]
    
    table2 = Table(table2_data, colWidths=[5*cm, 11*cm])
    table2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    
    story.append(table2)
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=Sujet_{request.exam_type}_{request.student_name or 'CANDIDAT'}.pdf"}
    )
