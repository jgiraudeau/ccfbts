from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models import Evaluation, User
from pydantic import BaseModel
from typing import List, Dict, Optional
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from fastapi.responses import StreamingResponse

router = APIRouter()

class ExportRequest(BaseModel):
    student_id: int
    exam_type: str # 'E4' or 'E6'
    evaluations: List[dict] # Simplified for now, passing raw data from frontend

# --- DOCX Generation ---
def create_docx(student_name: str, exam_type: str, data: List[dict]):
    document = Document()
    
    # Title
    title = document.add_heading(f'Bilan {exam_type} - {student_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f'Généré le {data[0].get("date", "Date inconnue")}')

    # Content
    for item in data:
        document.add_heading(f"Évaluation du {item.get('date')}", level=1)
        p = document.add_paragraph()
        p.add_run("Compétences évaluées :").bold = True
        
        # Table for skills
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Compétence'
        hdr_cells[1].text = 'Niveau'
        
        if 'ratings' in item:
            for skill_id, rating in item['ratings'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = skill_id # In real app, map ID to Name
                row_cells[1].text = rating
        
        document.add_paragraph().add_run("Commentaire :").bold = True
        document.add_paragraph(item.get('comment', 'Aucun commentaire.'))
        document.add_page_break()

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- PDF Generation ---
def create_pdf(student_name: str, exam_type: str, data: List[dict]):
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    elements.append(Paragraph(f"Bilan {exam_type} - {student_name}", styles['Title']))
    elements.append(Spacer(1, 12))

    for item in data:
        elements.append(Paragraph(f"Évaluation du {item.get('date', '')}", styles['Heading2']))
        
        # Table Data
        table_data = [['Compétence', 'Note']]
        if 'ratings' in item:
            for skill, rating in item['ratings'].items():
                table_data.append([skill, rating])
        
        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))
        
        elements.append(Paragraph("<b>Commentaire:</b>", styles['Normal']))
        elements.append(Paragraph(item.get('comment', 'Aucun commentaire.'), styles['Normal']))
        elements.append(Spacer(1, 24))

    doc.build(elements)
    file_stream.seek(0)
    return file_stream

@router.post("/export/docx")
def export_docx(request: ExportRequest, db: Session = Depends(get_db)):
    student = db.query(User).filter(User.id == request.student_id).first()
    student_name = student.name if student else "Etudiant"
    
    stream = create_docx(student_name, request.exam_type, request.evaluations)
    
    headers = {
        'Content-Disposition': f'attachment; filename="Bilan_{request.exam_type}_{student_name}.docx"'
    }
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

@router.post("/export/pdf")
def export_pdf(request: ExportRequest, db: Session = Depends(get_db)):
    student = db.query(User).filter(User.id == request.student_id).first()
    student_name = student.name if student else "Etudiant"
    
    stream = create_pdf(student_name, request.exam_type, request.evaluations)
    
    headers = {
        'Content-Disposition': f'attachment; filename="Bilan_{request.exam_type}_{student_name}.pdf"'
    }
    return StreamingResponse(stream, media_type="application/pdf", headers=headers)
